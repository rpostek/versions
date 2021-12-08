import yaml
import sys
import pathlib
import datetime
from docx import Document
from docx.shared import Pt
from pefile import PE
import PySimpleGUI as sg

class Settings:
    @staticmethod
    def read(filename):
        try:
            with open(filename, 'rt', encoding='utf8') as settings_file:
                return next(yaml.full_load_all(settings_file))
        except FileNotFoundError:
            #print('brak pliku konfiguracyjnego')
            sg.popup('brak pliku konfiguracyjnego')
            sys.exit(-1)
        except yaml.scanner.ScannerError:
            #print('błędy w pliku konfiguracyjnym')
            sg.popup('błędy pliku konfiguracyjnego')
            sys.exit(-2)
        except:
            #print("nieoczekiwany błąd przy wczytywaniu konfiguracji:", sys.exc_info()[0])
            sg.popup("nieoczekiwany błąd przy wczytywaniu konfiguracji")
            sys.exit(-3)


class Program:
    def __init__(self, nazwa, lokalizacja, producent=''):
        self.nazwa = nazwa
        self.lokalizacja = lokalizacja
        self.producent = producent
        self.data = ''
        self.wersja_produktu = ('', '', '', '')
        self.wersja_pliku = ('', '', '', '')
        self.wersja = ''
        self.plik = ''

    def __str__(self):
        return f'{self.nazwa}\t{self.producent}\t{self.lokalizacja}\t{self.data}\t{self.wersja}'

    def get_time(self):
        p = pathlib.Path(self.lokalizacja)
        self.data = datetime.date.fromtimestamp(p.stat().st_ctime).isoformat()
        self.plik = p.name

    def get_version(self):
        pe = PE(self.lokalizacja)
        if not 'VS_FIXEDFILEINFO' in pe.__dict__:
            return
        if not pe.VS_FIXEDFILEINFO:
            return
        verinfo = pe.VS_FIXEDFILEINFO[0]
        self.wersja_pliku = (verinfo.FileVersionMS >> 16, verinfo.FileVersionMS & 0xFFFF,
                             verinfo.FileVersionLS >> 16, verinfo.FileVersionLS & 0xFFFF)
        self.wersja_produktu = (verinfo.ProductVersionMS >> 16, verinfo.ProductVersionMS & 0xFFFF,
                                verinfo.ProductVersionLS >> 16, verinfo.ProductVersionLS & 0xFFFF)
        if self.wersja_produktu == self.wersja_pliku:
            self.wersja = ".".join(str(i) for i in self.wersja_pliku)
        else:
            self.wersja = str(self.wersja_produktu[0]) + '.' + ".".join(str(i) for i in self.wersja_pliku)

def update_document():
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)
    columns = config.get('dane', ['nazwa', 'producent', 'wersja'])
    table = document.add_table(rows=1, cols=len(columns), style=config.get('styl', 'Light List Accent 1'))
    hdr_cells = table.rows[0].cells
    for col in enumerate(columns, 0):
        hdr_cells[col[0]].text = col[1].replace('_', ' ')
    for p in aplikacje:
        row_cells = table.add_row().cells
        for col in enumerate(columns, 0):
            row_cells[col[0]].text = getattr(p, col[1])
    document.add_paragraph(f'Sporządzono {datetime.datetime.now().isoformat(sep=" ")[:19]}')

    path = pathlib.Path(config.get('dokument', '.'))
    filename = f"programy_{datetime.date.today().year}_{datetime.date.today().month:02}.docx"
    document.save(path.joinpath(filename))


if __name__ == '__main__':
    config = Settings.read('config.yaml')
    try:
        window = sg.Window("Wersje programów WBK 1.1", [[sg.Text("Zbieranie informacji:")],
              [sg.Listbox(values=[], size=(30,len(config['programy'])),
                        background_color='oldlace', text_color='black', no_scrollbar=True, key='-LIST-')]],
              finalize=True)
    except KeyError:
        sg.popup('błędy pliku konfiguracyjnego')
        sys.exit(-2)
    aplikacje = []
    for app in config['programy']:
        a = Program(app.get('nazwa', ''), app.get('lokalizacja', ''), app.get('producent', ''))
        window['-LIST-'].update(window['-LIST-'].get_list_values() + [a.nazwa,])
        event, values = window.read(timeout=100)
        if event == sg.WIN_CLOSED:
            window.close()
            sys.exit()
        try:
            a.get_time()
            a.get_version()
        except:
            a.wersja = 'błąd !!!'
        finally:
            #print(a.nazwa)
            aplikacje.append(a)
    update_document()
    window.close()